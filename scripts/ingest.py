#!/usr/bin/env python3
"""
LLM Wiki Ingest Pipeline — Main orchestrator.

Triggered by "学习入库" command from Lei Ge.
Full pipeline: file → dedup → extract → claims → conflict check → wiki write → report.

Usage:
    python ingest.py --file <path> [--file <path2> ...]
    python ingest.py --scan-temp
"""

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import uuid
from datetime import datetime

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)

import fitz
from docx import Document

from slug_generator import compute_sha256_prefix, generate_slug, extract_entities_from_content, detect_domains
from dedup import check_dedup, add_to_db
from multimodal_extract import assess_pdf_complexity, extract_with_rapidocr
from extract_claims import extract_claims_from_text, detect_claim_type
from conflict_resolver import resolve_conflict, build_notification_message

WIKI_PATH = os.environ.get('WIKI_PATH', '/home/agentuser/wiki')
TEMP_PATH = '/home/agentuser/temp'
RAW_TYPES = {
    '.txt': 'docs', '.md': 'docs', '.docx': 'docs', '.pdf': 'papers',
}

FRONTMATTER_TEMPLATE = """---
id: {id}
title: "{title}"
created: {created}
updated: {updated}
type: {page_type}
source: raw/{raw_subdir}/{raw_filename}
domain: [{domains}]
tags: [{tags}]
confidence: {confidence}
summary: "{summary}"
conflicts: [{conflicts}]
versions: [{versions}]
claims: {claims}
---

# {title}

## Summary
{summary}

## Details
{details}

## Related
- Sources: [[../raw/{raw_subdir}/{raw_filename}]]
{related_entities}
{related_concepts}

## Provenance
{provenance_block}
"""

PROVENANCE_TEMPLATE = """> [!source]| {claim_id}
> {claim_text}"""

def get_raw_subdir(file_path: str) -> str:
    return RAW_TYPES.get(os.path.splitext(file_path)[1].lower(), 'docs')

def extract_content(file_path: str) -> tuple[str, dict]:
    """Extract text content from file. Returns (text, metadata)."""
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    meta = {"method": "unknown", "text_density": 0.0}

    try:
        if ext in ('.txt', '.md'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            meta['method'] = 'direct'
        elif ext == '.docx':
            doc = Document(file_path)
            text = '\n'.join(p.text for p in doc.paragraphs)
            meta['method'] = 'python-docx'
        elif ext == '.pdf':
            assessment = assess_pdf_complexity(file_path)
            meta['text_density'] = assessment['text_density']
            meta['file_size_mb'] = assessment['file_size_mb']
            if assessment['recommendation'] == 'direct_ocr':
                text = extract_with_rapidocr(file_path)
                meta['method'] = 'rapidocr'
            elif assessment['recommendation'] == 'vlm':
                meta['method'] = 'vlm_needed'
                text = ""
            else:
                doc = fitz.open(file_path)
                # Skip cover page (page 0) if it has < 200 chars (likely a cover with fragmented text)
                start_page = 0
                if doc.page_count > 0:
                    first_page_chars = len(doc[0].get_text())
                    if first_page_chars < 200:
                        start_page = 1
                for page in doc[start_page:]:
                    text += page.get_text()
                doc.close()
                meta['method'] = 'fitz'
        elif ext == '':
            # Extensionless text files (e.g., Chinese news dumps from travelsky.com)
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            meta['method'] = 'direct'
        else:
            meta['method'] = 'unsupported'
    except Exception as e:
        meta['error'] = str(e)

    return text, meta

def detect_page_type(content: str, filename: str) -> str:
    content_lower = content.lower()
    if any(kw in content_lower for kw in ['会议', 'meeting', '纪要', '座谈']):
        return 'entity-event'
    if any(kw in content_lower for kw in ['获批', '入选', '荣获', '通过', '批准', '成立', '荣获', '获颁', '首次']):
        return 'entity-event'
    if any(kw in content_lower for kw in ['论文', 'paper', 'arxiv', '研究']):
        return 'entity-paper'
    if any(kw in content_lower for kw in ['公司', '企业', '集团', 'inc', 'ltd', 'co']):
        return 'entity-company'
    if any(kw in content_lower for kw in ['同志', '经理', '董事', '局长', '处长', '部长', 'ceo', 'cto']):
        return 'entity-person'
    return 'concept'

def estimate_confidence(content: str, page_type: str) -> float:
    """Estimate entity extraction confidence based on content quality.

    评分维度：
    1. 内容长度（越长实体越多）
    2. 实体关键词密度
    3. 政策文件特征（有"关于/意见/通知/办法"等格式的正文）
    4. 新闻报道特征（有发布时间/来源/浏览量元数据的，实体往往明确）
    """
    score = 0.5

    # Length bonus: longer content = more entities detectable
    if len(content) > 1000:
        score += 0.15
    elif len(content) > 500:
        score += 0.1
    elif len(content) > 200:
        score += 0.05

    # Entity signals in content
    entity_keywords = ['公司', '企业', '集团', '政策', '办法', '规定', '条例',
                       '人工智能', '航空', '民航', '数据', '模型', '系统']
    found = sum(1 for kw in entity_keywords if kw in content)
    score += min(found * 0.03, 0.15)

    # Page type clarity: event/company/person pages usually have clear facts
    if page_type in ('entity-event', 'entity-company', 'entity-person'):
        score += 0.15

    # Chinese text quality (no garbled chars)
    try:
        chinese_chars = len([c for c in content if '\u4e00' <= c <= '\u9fff'])
        if chinese_chars > 100:
            score += 0.1
        elif chinese_chars > 50:
            score += 0.05
    except Exception:
        pass

    # News format bonus: 发布时间/来源/浏览量 = factual reporting
    # These are event facts, not ambiguous policy analysis
    news_signals = sum(1 for sig in ['发布时间', '来源：', '浏览量', '作者：'] if sig in content)
    score += min(news_signals * 0.05, 0.10)

    return min(round(score, 2), 0.99)

def extract_summary(content: str) -> str:
    """Extract 3-sentence summary from content."""
    sentences = re.split(r'[。！？\n]', content.strip())
    sentences = [s.strip() for s in sentences if len(s.strip()) > 10][:3]
    if sentences:
        return '。'.join(sentences[:3]) + '。'
    return "内容摘要待补充。"

def build_provenance_block(claims: list, raw_filename: str) -> str:
    """Build markdown provenance section from claims."""
    if not claims:
        return "- No structured claims extracted."

    blocks = []
    for c in claims[:10]:  # Limit to 10 claims for readability
        claim_text = c['text'][:200].replace('\n', ' ').strip()
        blocks.append(PROVENANCE_TEMPLATE.format(
            claim_id=c['id'],
            claim_text=claim_text
        ))
    return '\n\n'.join(blocks)

def extract_claims_from_file_content(content: str, file_path: str) -> list[dict]:
    """Extract claims from content using extract_claims module."""
    try:
        return extract_claims_from_text(content, file_path)
    except Exception:
        # Fallback: treat whole content as one claim
        return [{
            "id": "c" + str(hash(content[:50]) % 100000),
            "text": content[:500],
            "para_index": 0,
            "entities": extract_entities_from_content(content[:500]),
            "domains": detect_domains(content[:500]),
            "type": detect_claim_type(content[:500], [], []),
            "source": file_path
        }]

def find_related_pages(content: str, wiki_path: str, page_type: str) -> list[str]:
    """Find existing wiki pages that might conflict or relate based on entity overlap."""
    entities = extract_entities_from_content(content[:1000])
    if not entities:
        return []

    # Search for existing pages with overlapping entities
    related = []
    search_dirs = [
        os.path.join(wiki_path, 'wiki', 'entities'),
        os.path.join(wiki_path, 'wiki', 'concepts'),
    ]

    for search_dir in search_dirs:
        if not os.path.exists(search_dir):
            continue
        for root, _, files in os.walk(search_dir):
            for f in files:
                if not f.endswith('.md'):
                    continue
                fpath = os.path.join(root, f)
                try:
                    with open(fpath, 'r', encoding='utf-8', errors='ignore') as fp:
                        page_content = fp.read()
                    page_entities = extract_entities_from_content(page_content[:500])
                    # If entities overlap, flag as potentially related
                    if set(entities) & set(page_entities):
                        rel_path = fpath.replace(wiki_path + '/', '')
                        related.append(rel_path)
                except:
                    continue

    return related[:5]  # Limit to 5

def find_existing_page_by_source(raw_source_path: str, wiki_path: str) -> str | None:
    """
    Find wiki page whose frontmatter 'source:' points to the given raw file path.
    Used for force_ingest to detect versioned document updates.
    raw_source_path: raw/{subdir}/{filename} format
    Returns: wiki page path relative to wiki root, or None if not found.
    """
    wiki_dir = os.path.join(wiki_path, 'wiki')
    if not os.path.exists(wiki_dir):
        return None

    for root, _, files in os.walk(wiki_dir):
        for f in files:
            if not f.endswith('.md') or f in ('index.md', 'log.md', '_redirects.yaml'):
                continue
            fpath = os.path.join(root, f)
            try:
                with open(fpath, 'r', encoding='utf-8', errors='ignore') as fp:
                    raw = fp.read()
            except:
                continue

            # Parse frontmatter
            fm_match = re.match(r'^---\n(.*?)\n---', raw, re.DOTALL)
            if not fm_match:
                continue
            fm_text = fm_match.group(1)
            source_match = re.search(r'^source:\s*(.+?)\s*$', fm_text, re.MULTILINE)
            if not source_match:
                continue
            page_source = source_match.group(1).strip()
            if page_source == raw_source_path or page_source.endswith('/' + os.path.basename(raw_source_path)):
                return fpath.replace(wiki_path + '/', '')

    return None

def update_existing_page_with_version(
    existing_page_rel_path: str,
    new_content: str,
    new_raw_subdir: str,
    new_raw_filename: str,
    new_claims: list,
    wiki_path: str
) -> dict:
    """
    Update an existing wiki page with new version content.
    - Bump updated date
    - Add old claim IDs to versions[]
    - Append conflict notice in Details
    - Add new claims to frontmatter
    Returns dict with update summary.
    """
    full_path = os.path.join(wiki_path, existing_page_rel_path)
    summary = {"updated": existing_page_rel_path, "versions_added": [], "claims_updated": False}

    try:
        with open(full_path, 'r', encoding='utf-8') as fp:
            content = fp.read()

        # Extract old claim IDs for version chain
        versions = []
        old_claim_ids = []
        claims_match = re.search(r'claims:\s*(\[.*?\])', content, re.DOTALL)
        if claims_match:
            try:
                old_claims = json.loads(claims_match.group(1))
                old_claim_ids = [c.get('id') for c in old_claims if c.get('id')]
            except:
                pass

        # Extract old source for versions chain
        source_match = re.search(r'^source:\s*(.+?)\s*$', content, re.MULTILINE)
        old_source = source_match.group(1).strip() if source_match else ""

        if old_source and old_source not in versions:
            versions.append(old_source)

        # Bump updated date
        content = re.sub(
            r'^updated:.*$',
            f'updated: {datetime.now().strftime("%Y-%m-%d")}',
            content, flags=re.MULTILINE
        )

        # Update source to new raw file
        new_source = f"raw/{new_raw_subdir}/{new_raw_filename}"
        if source_match:
            content = re.sub(
                r'^source:.*$',
                f'source: {new_source}',
                content, flags=re.MULTILINE
            )
        else:
            # Insert after first '---' of frontmatter
            content = content.replace('---\n', f'---\nsource: {new_source}\n', 1)

        # Append to versions chain
        versions_str = ', '.join(f'"{v}"' for v in versions)
        content = re.sub(
            r'versions:\s*\[.*?\]',
            f'versions: [{versions_str}]',
            content
        )

        # Update claims in frontmatter
        new_claims_json = json.dumps(new_claims, ensure_ascii=False, indent=2)
        content = re.sub(
            r'claims:\s*(\[.*?\])',
            f'claims: {new_claims_json}',
            content, flags=re.DOTALL
        )
        summary["claims_updated"] = True
        summary["versions_added"] = old_claim_ids

        # Append version notice in Details
        notice = (
            f'\n\n> [!version-update]\n'
            f'> **版本更新** by: `{new_raw_filename}`\n'
            f'> Previous source: `{old_source}`\n'
            f'> Claim IDs archived: `[{", ".join(str(c) for c in old_claim_ids[:5])}]`\n'
        )
        if '## Details' in content:
            content = content.replace('## Details', notice + '\n## Details', 1)
        else:
            content += notice

        with open(full_path, 'w', encoding='utf-8') as fp:
            fp.write(content)

        summary["status"] = "updated"

    except Exception as e:
        summary["status"] = "error"
        summary["error"] = str(e)

    return summary

def check_claim_conflicts(new_claims: list, existing_pages: list, wiki_path: str) -> list[dict]:
    """
    Check new claims against existing pages for conflicts.
    Returns list of conflict reports.
    """
    conflicts = []

    for new_claim in new_claims:
        new_type = new_claim.get('type', detect_claim_type(new_claim.get('text', ''), [], []))
        new_text = new_claim.get('text', '')

        # Skip analysis type — they don't conflict, they just coexist
        if new_type == 'analysis':
            continue

        for page_path in existing_pages:
            full_path = os.path.join(wiki_path, page_path) if not page_path.startswith('/') else page_path
            if not os.path.exists(full_path):
                continue

            try:
                with open(full_path, 'r', encoding='utf-8', errors='ignore') as fp:
                    page_content = fp.read()

                # Extract existing claims from page if any
                # Try to find claims in frontmatter or content
                existing_claims = []
                if 'claims:' in page_content:
                    # Claims are stored in frontmatter
                    cm = re.search(r'claims:\s*(\[.*?\])', page_content, re.DOTALL)
                    if cm:
                        try:
                            existing_claims = json.loads(cm.group(1))
                        except:
                            pass

                # Simple conflict detection: check if same entity + same type + different claim
                for old_claim in existing_claims:
                    old_type = old_claim.get('type', '')
                    old_text = old_claim.get('text', '')

                    # Only conflict if same type and entity overlap
                    if old_type == new_type and old_type in ('policy', 'data'):
                        # Check for textual similarity (simple keyword overlap)
                        old_words = set(re.findall(r'[\w]{3,}', old_text.lower()))
                        new_words = set(re.findall(r'[\w]{3,}', new_text.lower()))
                        overlap = old_words & new_words

                        if len(overlap) > 5:  # Significant overlap = same topic
                            # This is a potential conflict
                            resolution = resolve_conflict(old_claim, new_claim)
                            conflicts.append({
                                "old_claim": old_claim,
                                "new_claim": new_claim,
                                "page": page_path,
                                "resolution": resolution
                            })

            except Exception:
                continue

    return conflicts

def resolve_conflicts_and_update(conflicts: list, wiki_path: str, raw_filename: str) -> dict:
    """
    Process conflicts, apply auto-update or flag for human review.
    Returns summary dict.
    """
    summary = {
        "auto_updated": [],
        "flagged_human": [],
        "coexisting": [],
        "notifications": []
    }

    for conflict in conflicts:
        resolution = conflict['resolution']
        page = conflict['page']
        old_claim = conflict['old_claim']
        new_claim = conflict['new_claim']

        if resolution['resolution'] == 'auto_update':
            # Auto-update: append new version to versions[], update page
            try:
                full_path = os.path.join(wiki_path, page) if not page.startswith('/') else page
                with open(full_path, 'r', encoding='utf-8') as fp:
                    content = fp.read()

                # Add to versions
                versions_match = re.search(r'versions:\s*\[(.*?)\]', content)
                versions = []
                if versions_match:
                    try:
                        versions = json.loads('[' + versions_match.group(1) + ']')
                    except:
                        versions = [v.strip() for v in versions_match.group(1).split(',') if v.strip()]

                old_id = old_claim.get('id', 'unknown')
                if old_id not in versions:
                    versions.append(old_id)

                # Update versions line
                versions_items = ', '.join(f'"{v}"' for v in versions)
                content = re.sub(
                    r'versions:\s*\[.*?\]',
                    f'versions: [{versions_items}]',
                    content
                )

                # Update updated date
                content = re.sub(
                    r'^updated:.*$',
                    f'updated: {datetime.now().strftime("%Y-%m-%d")}',
                    content, flags=re.MULTILINE
                )

                # Add conflict notice to details
                notice = f'\n\n> [!conflict]\n> Updated by: {raw_filename} ({new_claim.get("id","new")})\n> Type: {resolution["type"]}\n> Action: Auto-updated per conflict resolution rules\n'
                if '## Details' in content:
                    content = content.replace('## Details', notice + '\n## Details', 1)

                with open(full_path, 'w', encoding='utf-8') as fp:
                    fp.write(content)

                summary['auto_updated'].append({
                    "page": page,
                    "old_claim": old_id,
                    "new_claim": new_claim.get('id'),
                    "type": resolution['type']
                })

                # Build notification
                notif = build_notification_message(resolution, page)
                summary['notifications'].append(notif)

            except Exception as e:
                summary['flagged_human'].append({
                    "page": page,
                    "error": str(e),
                    "resolution": resolution
                })

        elif resolution['resolution'] == 'flag_human':
            summary['flagged_human'].append({
                "page": page,
                "old_claim": old_claim,
                "new_claim": new_claim,
                "resolution": resolution
            })
            notif = build_notification_message(resolution, page)
            summary['notifications'].append(notif)

        else:  # keep_both
            summary['coexisting'].append({
                "page": page,
                "new_claim": new_claim.get('id')
            })
            # No notification needed — different perspectives

    return summary

def write_raw_file(file_path: str, wiki_raw_dir: str) -> tuple:
    """Copy file to raw/ with English name. Returns (subdir, filename, dest_path)."""
    content_preview = ""
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == '.pdf':
            import fitz
            doc = fitz.open(file_path)
            for page in doc[:3]:
                content_preview += page.get_text()
            doc.close()
        elif ext in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content_preview = f.read(2000)
        elif ext == '.docx':
            from docx import Document
            doc = Document(file_path)
            for para in doc.paragraphs[:20]:
                content_preview += para.text + '\n'
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content_preview = f.read(500)
    except Exception:
        pass

    slug_result = generate_slug(file_path, content_preview)
    sha_prefix = compute_sha256_prefix(file_path)
    new_filename = f"{sha_prefix}_{slug_result}{ext}"
    subdir = get_raw_subdir(file_path)
    dest_dir = os.path.join(wiki_raw_dir, subdir)
    os.makedirs(dest_dir, exist_ok=True)
    shutil.copy2(file_path, os.path.join(dest_dir, new_filename))
    return subdir, new_filename, os.path.join(dest_dir, new_filename)

def build_wiki_page_content(content: str, raw_filename: str, raw_subdir: str,
                              page_type: str, slug: str, domains: list,
                              claims: list, conflicts_resolved: dict,
                              estimated_confidence: float = 0.85) -> str:
    """Build full wiki page markdown with frontmatter and provenance."""
    # title: CamelCase English, no separators, no SHA prefix
    # e.g. "sha_travelsky-jiaolei-ai-aviation" → "TravelskyJiaoleiAiAviation"
    slug_no_sha = re.sub(r'^[a-f0-9]{8}_', '', slug)  # strip SHA prefix
    title = slug_no_sha.replace('-', ' ').replace('_', ' ').title().replace(' ', '')
    summary = extract_summary(content)

    entities = extract_entities_from_content(content[:1000])
    detected_domains = detect_domains(content[:1000])
    # Wikilink format: [[entity_xxx]] and [[concept_xxx]] (underscore, NOT hyphen)
    entity_links = '\n'.join([f'- Entities: [[entity_{e}]]' for e in entities[:5]])
    concept_links = '\n'.join([f'- Concepts: [[concept_{d}]]' for d in detected_domains[:5]])
    # Fix: convert hyphens to underscores in wikilinks (SKILL.md: wikilinks use underscore NOT hyphen)
    # e.g. [[entity_strategic-emerging-industry]] → [[entity_strategic_emerging_industry]]
    entity_links = re.sub(r'\[\[entity_([a-z]+(?:[-][a-z]+)+)', lambda m: f'[[entity_{m.group(1).replace("-", "_")}', entity_links)
    concept_links = re.sub(r'\[\[concept_([a-z]+(?:[-][a-z]+)+)', lambda m: f'[[concept_{m.group(1).replace("-", "_")}', concept_links)

    # Build claims JSON for frontmatter
    claims_json = json.dumps(claims, ensure_ascii=False)

    # Build provenance block
    provenance_block = build_provenance_block(claims, raw_filename)

    # Check if there are auto-updated conflicts for this page
    auto_updates = [u for u in conflicts_resolved.get('auto_updated', [])]
    conflicts_str = json.dumps([u['old_claim'] for u in auto_updates]) if auto_updates else ""

    versions_list = []
    for u in auto_updates:
        if u['old_claim'] not in versions_list:
            versions_list.append(u['old_claim'])
    versions_str = ', '.join(f'"{v}"' for v in versions_list) if versions_list else ""

    details = content[:2000]
    if len(content) > 2000:
        details += '\n\n[... truncated ...]\n\n' + content[-500:]

    domains_str = ', '.join(domains) if domains else 'general'
    tags_str = page_type.replace('entity-', '') + ', ' + domains_str

    return FRONTMATTER_TEMPLATE.format(
        id=str(uuid.uuid4())[:8],
        created=datetime.now().strftime('%Y-%m-%d'),
        updated=datetime.now().strftime('%Y-%m-%d'),
        page_type=page_type,
        raw_subdir=raw_subdir,
        raw_filename=raw_filename,
        domains=', '.join(f'"{d}"' for d in domains) if domains else '"general"',
        tags=f'"{tags_str}"',
        confidence=str(estimated_confidence),
        summary=summary.replace('"', '\\"'),
        conflicts=conflicts_str if conflicts_str else "",
        versions=versions_str if versions_str else "",
        claims=claims_json,
        title=title,
        details=details,
        related_entities=entity_links if entity_links else "",
        related_concepts=concept_links if concept_links else "",
        provenance_block=provenance_block
    )

def ingest_file(file_path: str, wiki_path: str = WIKI_PATH) -> dict:
    """Full ingest pipeline for one file."""
    result = {
        "file": file_path,
        "status": "unknown",
        "raw_path": None,
        "wiki_pages": [],
        "review_needed": False,
        "dedup_status": None,
        "conflicts": {},
        "notifications": [],
        "error": None,
        "claims_types": {}
    }

    try:
        # Step 1: Dedup
        dedup_result = check_dedup(file_path)
        result['dedup_status'] = dedup_result['status']
        result['dedup_similarity'] = dedup_result.get('similarity', 0.0)
        is_force_ingest = (dedup_result['action'] == 'force_ingest')
        if dedup_result['action'] == 'skip':
            result['status'] = 'skipped'
            result['reason'] = f"{dedup_result['status']} (similarity: {result['dedup_similarity']:.2f})"
            return result

        # Step 2: Extract content
        content, extract_meta = extract_content(file_path)
        result['extraction_method'] = extract_meta.get('method', 'unknown')
        if not content or len(content.strip()) < 50:
            result['status'] = 'error'
            result['error'] = 'Content too short or empty'
            return result

        # Step 3: Extract claims (for provenance)
        claims = extract_claims_from_file_content(content, file_path)
        result['claims_extracted'] = len(claims)
        result['claims_types'] = {c['type']: result['claims_types'].get(c['type'], 0) + 1 for c in claims}

        # Step 4: Write to raw/
        wiki_raw_dir = os.path.join(wiki_path, 'raw')
        os.makedirs(wiki_raw_dir, exist_ok=True)
        raw_subdir, new_filename, raw_dest = write_raw_file(file_path, wiki_raw_dir)
        result['raw_path'] = f"raw/{raw_subdir}/{new_filename}"

        # Step 5: Add to dedup DB
        add_to_db(file_path, wiki_filepath=raw_dest)

        # Step 6: Detect page type + domains
        page_type = detect_page_type(content, os.path.basename(file_path))
        domains = detect_domains(content[:1000])
        estimated_confidence = estimate_confidence(content, page_type)

        # Step 7: Find related pages for conflict check
        related_pages = find_related_pages(content, wiki_path, page_type)

        # Step 8: Check for claim conflicts
        conflicts = check_claim_conflicts(claims, related_pages, wiki_path)
        result['conflicts_checked'] = len(conflicts)

        # Step 9: Resolve conflicts
        conflicts_resolved = {"auto_updated": [], "flagged_human": [], "coexisting": [], "notifications": []}
        if conflicts:
            conflicts_resolved = resolve_conflicts_and_update(conflicts, wiki_path, new_filename)
            result['conflicts'] = conflicts_resolved
            result['notifications'] = conflicts_resolved.get('notifications', [])

        # Step 10: Build wiki page content with provenance
        page_slug = new_filename.rsplit('.', 1)[0]
        page_content = build_wiki_page_content(
            content, new_filename, raw_subdir,
            page_type, page_slug, domains,
            claims, conflicts_resolved,
            estimated_confidence
        )

        # Step 11: Write wiki page
        page_slug = new_filename.rsplit('.', 1)[0]
        if 'entity-' in page_type:
            entity_subtype = page_type.replace('entity-', '')
            page_dir = os.path.join(wiki_path, 'wiki/entities', entity_subtype)
        else:
            page_dir = os.path.join(wiki_path, 'wiki', page_type + 's')
        os.makedirs(page_dir, exist_ok=True)

        page_path = os.path.join(page_dir, f"{page_slug}.md")

        # If auto-updated via conflict resolution, don't overwrite the existing page
        auto_updated_slugs = [u['page'].replace('wiki/', '').replace('.md', '') for u in conflicts_resolved.get('auto_updated', [])]
        if page_slug in auto_updated_slugs:
            result['status'] = 'merged'
            result['wiki_pages'] = [u['page'] for u in conflicts_resolved.get('auto_updated', []) if page_slug in u['page']]
        elif is_force_ingest:
            # Force-ingest: check if this raw file corresponds to an existing page (version update)
            raw_source_key = f"raw/{raw_subdir}/{new_filename}"
            existing_page_rel = find_existing_page_by_source(raw_source_key, wiki_path)
            if existing_page_rel:
                # Update existing page with new version content
                version_update = update_existing_page_with_version(
                    existing_page_rel, content, raw_subdir, new_filename, claims, wiki_path
                )
                result['status'] = version_update.get('status', 'version_updated')
                result['wiki_pages'] = [existing_page_rel]
                result['version_update'] = version_update
            else:
                # No existing page found — treat as normal new page (edge case)
                with open(page_path, 'w', encoding='utf-8') as f:
                    f.write(page_content)
                result['wiki_pages'].append(page_path)
                result['status'] = 'success'
        else:
            with open(page_path, 'w', encoding='utf-8') as f:
                f.write(page_content)
            result['wiki_pages'].append(page_path)
            result['status'] = 'success'

        result['page_type'] = page_type
        result['domains'] = domains

        # Step 12: Check confidence → review (use pre-computed estimated_confidence)
        review_triggered = (
            'unknown' in page_slug
            or conflicts_resolved.get('flagged_human')
            or estimated_confidence < 0.9
        )
        if review_triggered:
            result['review_needed'] = True
            result['status'] = 'review_needed'
            # Actually move the file to review/
            review_dir = os.path.join(wiki_path, 'wiki/review')
            os.makedirs(review_dir, exist_ok=True)
            review_path = os.path.join(review_dir, os.path.basename(page_path))
            if os.path.exists(page_path):
                shutil.move(page_path, review_path)
                result['wiki_pages'] = [review_path]
                result['review_reason'] = (
                    f"confidence={estimated_confidence}" if estimated_confidence < 0.9
                    else ("slug='unknown'" if 'unknown' in page_slug else "human_flagged")
                )

    except Exception as e:
        result['status'] = 'error'
        result['error'] = str(e)

    return result

def scan_temp() -> dict:
    if not os.path.exists(TEMP_PATH):
        return {"files": [], "message": "temp/ does not exist"}
    files = []
    for f in os.listdir(TEMP_PATH):
        full = os.path.join(TEMP_PATH, f)
        if os.path.isfile(full):
            files.append({"name": f, "path": full, "size_bytes": os.path.getsize(full),
                          "ext": os.path.splitext(f)[1].lower()})
    return {"files": files, "count": len(files)}


def auto_git_commit(wiki_path: str, ingest_count: int) -> dict:
    """
    自动 git add + commit（触发 post-commit hook QA）。
    仅在有变更时 commit。
    """
    try:
        # Check if there are changes
        result = subprocess.run(
            ['git', 'status', '--porcelain'],
            cwd=wiki_path,
            capture_output=True, text=True, timeout=30
        )
        if not result.stdout.strip():
            return {"status": "no_changes", "message": "无变更，跳过 commit"}

        # git add .
        subprocess.run(['git', 'add', '.'], cwd=wiki_path, capture_output=True, timeout=30)

        # git commit with message
        today = datetime.now().strftime('%Y-%m-%d')
        msg = f"ingest: {ingest_count} file(s) ({today})"
        result = subprocess.run(
            ['git', 'commit', '-m', msg],
            cwd=wiki_path,
            capture_output=True, text=True, timeout=30
        )

        if result.returncode == 0:
            return {"status": "committed", "message": msg}
        else:
            return {"status": "error", "message": result.stderr}
    except Exception as e:
        return {"status": "error", "message": str(e)}


def main():
    parser = argparse.ArgumentParser(description='LLM Wiki Ingest Pipeline')
    parser.add_argument('--file', action='append', help='File to ingest')
    parser.add_argument('--scan-temp', action='store_true', help='Scan temp/')
    parser.add_argument('--wiki-path', default=WIKI_PATH, help='Wiki root')
    args = parser.parse_args()

    if args.scan_temp:
        print(json.dumps(scan_temp(), ensure_ascii=False, indent=2))
        return

    if not args.file:
        parser.print_help()
        return

    results = []
    all_notifications = []
    for fpath in args.file:
        if not os.path.exists(fpath):
            results.append({"file": fpath, "status": "error", "error": "File not found"})
            continue
        r = ingest_file(fpath, args.wiki_path)
        if r.get('notifications'):
            all_notifications.extend(r['notifications'])
        results.append(r)

    output = {"ingest_results": results}
    if all_notifications:
        output["notifications"] = all_notifications

    print(json.dumps(output, ensure_ascii=False, indent=2))

    # 自动 git commit（触发 post-commit hook QA）
    ingest_count = len([r for r in results if r.get('status') == 'success'])
    commit_result = auto_git_commit(args.wiki_path, ingest_count)
    print(f"\n[git] {commit_result['message']}")

if __name__ == '__main__':
    main()
